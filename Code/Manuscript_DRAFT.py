# -*- coding: utf-8 -*-
"""
Created on Tue Feb 17 10:50:00 2015

@author: Alex
"""
plt.ioff()
from docx import *
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
##  Create Document
document = Document()

######## SOME TOOLS
def add_figure_caption(fig_num=str(len(document.inline_shapes)),caption=''):
    cap = document.add_paragraph("Figure "+fig_num+". "+caption)
    cap.paragraph_style = 'caption'
    return
    
def dataframe_to_table(df=pd.DataFrame(),table_num=str(len(document.tables)+1),caption=''):
    table = document.add_table(rows=1, cols=len(df.columns)) 
    ## Merge all cells in top row and add caption text
    table_caption = table.rows[0].cells[0].merge(table.rows[0].cells[len(df.columns)-1])
    table_caption.text = "Table "+table_num+". "+caption
    ## Add  header
    header_row = table.add_row().cells
    col_count =0 ## counter  to iterate over columns
    for col in  header_row:
        col.text = df.columns[col_count] #df.columns[0] is the index
        col_count+=1
    ## Add data by  iterating over the DataFrame rows, then using a dictionary of DataFrame column labels to extract data
    col_labels = dict(zip(range(len(df.columns)),df.columns.tolist())) ## create dictionary where '1  to  n' is key for DataFrame columns
    for row in df.iterrows():  ## iterate over  rows in  DataFrame
        #print row[1]
        row_cells = table.add_row().cells ## Add a row to the  table
        col_count  = 0
        for cell in row_cells: ## iterate over the columns in the row
            #print row[1][str(col_labels[col_count])]
            cell.text = str(row[1][str(col_labels[col_count])]) ## and plug in data using a dictionary to  get column labels for DataFrame
            col_count+=1
    ## Format Table Style  
    table.style = 'TableGrid'   
    table.autofit
    #table.num = str(len(document.tables)+1)
    return table
    
def add_equation(eq_table):
    t = document.add_table(rows=len(eq_table.rows),cols=len(eq_table.columns))
    t.rows[0].cells[1].text = eq_table.rows[0].cells[1].text
    t.rows[0].cells[2].text = eq_table.rows[0].cells[2].text   
    if len(eq_table.rows)>1:
        t.rows[1].cells[0].merge(t.rows[1].cells[2]).text = eq_table.rows[1].cells[0].text
    t.style = 'TableGrid'   
    t.autofit
    return t
###################################################################################################################################################################    


#### TABLES ########################################################################################################################################################
### Landcover_Table
table_count=0
def tab_count():
    global table_count
    table_count+=1
    return str(table_count)
# Prepare LULC Data
landcover_table = LandCover_table()
landcover_table.table_num = str(tab_count())

### Storm Water Discharge Table
Q_Diff_table = Q_storm_diff_table() ## function to create table data
Q_Diff_table.table_num = str(tab_count())

### Storm Sediment Discharge Table
## Prepare table data
S_Diff_table = S_storm_diff_table() ## function to create table data
S_Diff_table.table_num = str(tab_count())

S_Diff_table_quarry = S_storm_diff_table_quarry()
S_Diff_table_quarry.table_num = str(tab_count())

### Storm Q and SSY summary table
Q_S_Diff_summary_table = Q_S_storm_diff_summary_table()
Q_S_Diff_summary_table.table_num = str(tab_count())

### Pearson r coefficient table
Pearson_table = Pearson_r_Table()
Pearson_table.table_num = str(tab_count())

### Pearson r coefficient table
Spearman_table = Spearman_r_Table()
Spearman_table.table_num = str(tab_count())

#### FIGURES ########################################################################################################################################################
figure_count=0
def fig_count():
    global figure_count
    figure_count+=1
    return str(figure_count)
### INTRODUCTION
## Study Area Map
Study_Area_map = {'filename':maindir+'Figures/Maps/FagaaluInstruments land only map.tif', 'fig_num':str(fig_count())}
## Quarry_picture
Quarry_picture = {'filename':maindir+'Figures/Maps/Quarry before and after.tif','fig_num':str(fig_count())}

### RESULTS
### Stage-Discharge Rating Curves
## DAM_StageDischarge
DAM_StageDischarge = {'filename':figdir+"Q/Water Discharge Ratings for FOREST (DAM)",'fig_num':str(fig_count())}
plotQratingDAM(ms=8,show=False,log=False,save=True,filename=DAM_StageDischarge['filename'])
## LBJ_StageDischarge
LBJ_StageDischarge = {'filename':figdir+"Q/Water Discharge Ratings for VILLAGE (LBJ)",'fig_num':str(fig_count())}
plotQratingLBJ(ms=8,show=False,log=False,save=True,filename=LBJ_StageDischarge['filename'])

### SSC
## SSC Boxplots
SSC_Boxplots = {'filename':figdir+'SSC/Grab sample boxplots','fig_num':str(fig_count())}
plotSSCboxplots(storm_samples_only=True,withR2=False,show=False,save=True,filename=SSC_Boxplots['filename'])
## Discharge vs Sediment Concentration
Discharge_Concentration = {'filename':figdir+'SSC/Water discharge vs Sediment concentration','fig_num':str(fig_count())}
plotQvsC(storm_samples_only=False,ms=6,show=False,log=False,save=True,filename=Discharge_Concentration['filename'])

### T-SSC Rating Curves
## Synthetic Rating Curves
Synthetic_Rating_Curve = {'filename':figdir+'T/Synthetic Rating Curves Fagaalu','fig_num':str(fig_count())} ## define file name to find the png file from other script
Synthetic_Rating_Curves_Fagaalu(param='SS_Mean',show=False,save=True,filename=Synthetic_Rating_Curve['filename'])## generate figure from separate script and save to file
## LBJ and DAM YSI T-SSC rating curves
LBJ_and_DAM_YSI_Rating_Curve = {'filename':figdir+'T/T-SSC rating LBJ and DAM YSI','fig_num':str(fig_count())}
plotYSI_compare_ratings(DAM_YSI,DAM_SRC,LBJ_YSI,Use_All_SSC=False,show=False,save=True,filename=LBJ_and_DAM_YSI_Rating_Curve['filename'])
## LBJ OBSa T-SSC rating curve
LBJ_OBSa_Rating_Curve = {'filename':figdir+'T/T-SSC rating LJB OBSa','fig_num':str(fig_count())}
OBSa_compare_ratings(df=LBJ_OBSa,df_SRC=LBJ_SRC,SSC_loc='LBJ',Use_All_SSC=False,show=False,save=True,filename=LBJ_OBSa_Rating_Curve['filename'])  
## LBJ OBSa T-SSC rating curve
LBJ_OBSb_Rating_Curve = {'filename':figdir+'T/T-SSC rating LJB OBSb','fig_num':str(fig_count())}
OBSb_compare_ratings(df=LBJ_OBSb,df_SRC=LBJ_SRC,SSC_loc='LBJ',Use_All_SSC=False,show=False,save=True,filename=LBJ_OBSb_Rating_Curve['filename'])  

### Storms
Example_Storm = {'filename':figdir+'storm_figures/Example_Storm','fig_num':str(fig_count())}
plot_storm_individually(LBJ_storm_threshold,LBJ_StormIntervals.loc[63],show=False,save=True,filename=Example_Storm['filename']) 

### SSY models
SSY_models_ALL = {'filename':figdir+'SSY/SSY Models ALL','fig_num':str(fig_count())}
plotALLStorms_ALLRatings(subset='pre',ms=4,norm=True,log=True,show=False,save=True,filename=SSY_models_ALL['filename'])

###### EQUATIONS ############################################################################################################################################
equation_count=0
def eq_count():
    global equation_count
    equation_count+=1
    return str(equation_count)
    
Equations = Document(maindir+'/Manuscript/Equations.docx').tables
## SSYev = Q*SSC 
SSYEV_eq = Equations[0].table
SSYEV_eq.eq_num = eq_count()
## DR = SSY/SSYpre
DR_eq = Equations[2].table
DR_eq.eq_num = eq_count()
## SSYpre = Area_total * (SSYupper/AREAupper)
SSYpre_eq = Equations[3].table
SSYpre_eq.eq_num = eq_count()
## predict_SSYev = aXb
predict_SSYEV_eq = Equations[1].table
predict_SSYEV_eq.eq_num = eq_count()
## PE = sqrt(sum(Error^2+Error^2))
PE_eq = Equations[4].table
PE_eq.eq_num = eq_count()
############################################################################################################################################

#### TITLE
title_title = document.add_heading('TITLE:',level=1)
title_title.paragraph_format.space_before = 0
title = document.add_heading('Contributions of human activities to suspended sediment yield during storm events from a steep, small, tropical watershed',level=1)
title.paragraph_format.space_before = 0
#### ABSTRACT
abstract_title = document.add_heading('ABSTRACT',level=2)
abstract_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract = document.add_paragraph('Abstract text goes here....')
### Read in Body Text
Body_Text = Document(maindir+'/Manuscript/Body_Text.docx')
section_count = 0
Section = {'Introduction':[],'Study Area':[],'Methods':[],'Results':[],'Discussion':[],'Conclusion':[]}
for paragraph in Body_Text.paragraphs:
    if len(paragraph.text)==0:
        section_count+=1
    if section_count == 1:
        Section['Introduction'].append(paragraph)
    if section_count == 2:
        Section['Study Area'].append(paragraph)   
    if section_count == 3:
        Section['Methods'].append(paragraph)
    if section_count == 4:
        Section['Results'].append(paragraph)        
    if section_count == 5:
        Section['Discussion'].append(paragraph)        
    if section_count == 6:
        Section['Conclusion'].append(paragraph)    
#### INTRODUCTION
introduction_title = document.add_heading('Introduction',level=2)
for paragraph in Section['Introduction'][2:]:
    p = document.add_paragraph(paragraph.text)
    p.style='BodyText'
    p.paragraph_format.left_indent = 0
    p.paragraph_format.first_line_indent = Inches(.4)

#### STUDY AREA
study_area_title = document.add_heading('Study Area',level=2)
## Study Area map
if 'Study_Area_map' in locals():
    document.add_picture(Study_Area_map['filename'],width=Inches(6))
    add_figure_caption(Study_Area_map['fig_num'],"Faga'alu watershed showing upper (undisturbed) and lower (human-disturbed) subwatersheds.")
# Climate
document.add_heading('Climate',level=3)
document.add_paragraph("Precipitation on Tutuila (14S, 170W) is caused by major storms including cyclones and tropical depressions, isolated thunderstorms, and orographic uplifting of trade-wind squalls over the high (300-600m), mountainous ridge that runs the length of the island. Unlike many other Pacific Islands, the mountainous ridge runs parallel to the predominant wind direction, and does not cause a significant windward/leeward rainfall gradient. When controlling for drainage area, average annual specific discharge (m3/yr/km2) shows little spatial variation across the island, irrespective of location or orientation (Dames & Moore, 1981). Precipitation varies orographically from an average 6,350mm/yr at high elevation to 2,380mm/yr at the shoreline, averaging 3,800mm/yr over the island from 1903 to 1973 (Eyre, 1994). In Faga'alu watershed, rainfall records show average precipitation is 6,350 mm at Matafao Mtn. (653m m.a.s.l), 5,280mm at Matafao Reservoir (249m m.a.s.l.) and about 3,800mm on the coastal plain (Craig, 2009; Dames & Moore, 1981; Tonkin & Taylor International Ltd., 1989; Wong, 1996). Potential evapotranspiration follows the opposite trend, with annual mean PET varying from 890mm at high elevation to 1,150mm at sea level (Izuka, 2005). Tropical cyclones are erratic but recently have occurred on average every 1-13 years (Craig, 2009) and bring intense rainfall, flooding, landslides, and high sediment loading events (Buchanan-Banks, 1979). ")
document.add_paragraph("There are two subtle rainfall seasons: a drier winter season, from June to September (32% of annual P) and a wetter summer season, from October to May (68% of annual P). During the drier winter season, the island is influenced by the southeast Trades and relatively stronger, predominantly East to Southeast winds, lower temperatures, lower humidity and lower total rainfall. During the wetter summer season the Inter-Tropical Convergence Zone (ITCZ) moves over the island, causing light to moderate Northerly winds, higher temperatures, higher humidity, and higher total rainfall. While total rainfall is lower in the drier Tradewind season, large peak rainfalls are still observed. Analysis of 212 peak discharges at 11 continuous-record gaging sites up to 1990 showed 65.5% of annual peak flows occurred during the summer wet season (Wong, 1996). ")
# Land Use
document.add_heading('Land Use',level=3)
document.add_paragraph("Faga'alu watershed (1.86 km2) is characterized by large areas of undisturbed, steeply sloping, heavily forested hillsides in the upper watershed, and similarly steep topography with relatively small flat areas that are urbanized or densely settled in the lower watershed (Figure "+Study_Area_map['fig_num']+"). This settlement pattern is typical for volcanic islands with steep topography in the south Pacific. Monitoring efforts focused on Faga'alu Stream, which discharges to a sediment-impacted reef (Aeby et al., 2006), and includes two unique features not found in "+'"typical"'+" watersheds in American Samoa: 1) an open aggregate quarry (QUARRY, Figure "+Study_Area_map['fig_num']+"), and 2) a large impervious area associated with a hospital. Three water impoundment structures were built in the upper Faga'alu watershed for drinking water supply and hydropower but only the highest, Matafao Reservoir, was ever connected to the municipal water system and has since fallen out of use (Tonkin & Taylor International Ltd., 1989)(Figure "+Study_Area_map['fig_num']+").")
document.add_paragraph("Land use in Faga'alu watershed includes agriculture, roads, and urbanization (Table "+landcover_table.table_num+"), but the predominant land cover is undisturbed forest on the steep hillsides (85.7%). These forests are prone to natural landslides that can contribute large amounts of sediment during storm events (Buchanan-Banks, 1979; Calhoun and Fletcher, 1999). Compared to other watersheds on Tutuila, a relatively large portion of Faga'alu watershed is urbanized ("+'"high intensity developed"'+" in Table "+landcover_table.table_num+", 4.6%), due to large areas of impervious surface associated with the hospital and the numerous residences and businesses. A small portion of the watershed (1.1%) is developed open space, which includes landscaped lawns and parks. In addition to some small, household gardens there are several small agricultural areas growing banana and taro on the steep hillsides. NOAA's Land Cover map (2.5m res.) classified the agricultural plots as "+'"Grassland"'+" due to the high grass cover in the plots (Table "+landcover_table.table_num+") (NOAA's Ocean Service and Coastal Services Center, 2010). These plots are currently receiving technical assistance from the Natural Resource Conservation Service (NRCS) to mitigate erosion problems.")
## Land Use/Land cover Table
if 'landcover_table' in locals():
    dataframe_to_table(df=landcover_table,table_num=landcover_table.table_num,caption="Land use categories in Fag'alu and Nu'uuli watersheds CITATION")
document.add_paragraph("")
## Quarry description
document.add_paragraph("In Faga'alu there is an open-pit aggregate quarry (~2ha) that accounts for the majority of the 1.1% bare land area in Faga'alu watershed (Table "+landcover_table.table_num+"). The quarry has been in continuous operation since the 1960's by advancing into the steep hillside to quarry the underlying basalt formation (Latinis 1996). The quarry operators have installed some sediment management practices such as silt fences and settling ponds (Horsley-Witten, 2011) but they are unmaintained and inadequate to control the large amount of sediment mobilized by the intense tropical rains (Horsley-Witten, 2012a). Longitudinal sampling of Faga'alu stream in 2011 showed significantly increased turbidity downstream of the quarry and of a new bridge construction site on the village road (Curtis et al., 2011). Construction of the bridge was completed March 2012 and no longer increases turbidity. There are several small footpaths and unpaved driveways, but most unpaved roads are stabilized with compacted gravel and do not appear to be a major contributor of sediment (Horsley-Witten, 2012b).")
## Quarry picture
if 'Quarry_picture' in locals():
    document.add_picture(Quarry_picture['filename'],width=Inches(6))
    add_figure_caption(Quarry_picture['fig_num'],"Photos of the open aggregate quarry in Faga'alu in 2012 (Top) and 2014 (Bottom). Photo: Messina")

#### METHODS
methods_title = document.add_heading('Methods',level=2)
document.add_paragraph("A nested-watershed approach was used to quantify the contribution of SSY from undisturbed and human-disturbed areas to total SSY to Faga'alu Bay during baseflow, and during storm events of varying magnitude (SSYEV). While steep, mountainous streams can discharge large amounts of bedload (Milliman and Syvitski, 1992), this research is focused on sediment size fractions that can be transported in suspension in the marine environment to settle on corals, and this is generally restricted to silt and clay fractions (<16um) (Asselman, 2000).")
document.add_paragraph("SSY generated by individual storm events (SSYEV) can be used to assess the contribution of individual subwatersheds to total SSY (Zimmermann et al., 2012), compare the responses of different watersheds (Basher et al., 2011; Duvert et al., 2012; Fahey et al., 2003; Hicks, 1990),  and determine changes in SSY from the same watershed over time (Bonta, 2000). SSYEV is calculated by integrating continuous suspended sediment load from measured or modeled discharge (Q) and suspended sediment concentration (SSC) (Duvert et al., 2012):")
add_equation(SSYEV_eq) ## Equation
### Disturbance Ratio
document.add_heading('Disturbance Ratio',level=3)
document.add_paragraph("To assess the contribution of each subwatershed to total SSY from the watershed, the percent contribution of each subwatershed was calculated as the difference between SSYEV observed at the upstream and downstream monitoring stations(FOREST, VILLAGE, Figure "+Study_Area_map['fig_num']+"). Another approach is the disturbance ratio (DR), which is the ratio of SSYEV from the total human-disturbed watershed under current conditions (measured at the watershed outlet (VILLAGE)) to SSY under pre-disturbance conditions (SSYpre):")
add_equation(DR_eq) ## Equation
document.add_paragraph("SSYpre is calculated assuming that the specific SSY (Mg/km2) from forested parts of the lower watershed is similar to the specific SSY from the upper watershed:")
add_equation(SSYpre_eq) ## Equation
document.add_paragraph("The percent contribution and DR were calculated for each storm event and averaged to determine the contribution of the Upper and Lower subwatersheds to total SSY.") ###TODO DR for each storm?

### Predicting event suspended sediment yield (SSYEV)
document.add_heading("Predicting event suspended sediment yield (SSYEV)",level=3)
document.add_paragraph("SSYEV may be correlated with precipitation or discharge variables ("+'"storm metrics"'+"), so four storm metrics were tested in this research: total event precipitation (Psum), event rainfall erosivity (EI30) (Hicks, 1990), total event water discharge (Qsum), and peak event water discharge (Qmax) (Duvert et al., 2012; Rodrigues et al., 2013). Storm metrics may be linearly or nonlinearly correlated with SSYEV, so both Pearson's and Spearman's correlation coefficients were calculated to select the best predictor of SSYEV from the total watershed, and from each subwatershed.")
document.add_paragraph("The relationship between SSYEV and storm metrics may be a linear function, but is often best fit by a watershed-specific power law function of the form:")
add_equation(predict_SSYEV_eq) ## Equation

### Data Collection
document.add_heading('Data Collection',level=3)
document.add_heading('Precipitation',level=4)
document.add_paragraph("Precipitation (P) was measured with Rainwise RAINEW tipping-bucket rain gages at 1 min intervals at three locations in Faga'alu watershed (Figure "+Study_Area_map['fig_num']+"). Data at RG2 was only recorded January-March, 2012, to determine a relationship between elevation and precipitation. Precipitation at 15 min intervals was also measured at the Vantage Pro Weather Station (Wx) and used to fill any data gaps in the precipitation recorded at RG1. The total event precipitation (Psum) and event rainfall erosivity (EI30) were calculated using data from RG1, with data gaps filled by data from Wx.")
## Water Discharge
document.add_heading('Water Discharge',level=4)
document.add_paragraph("Water discharge (Q) in Faga'alu Stream was derived from 15 min interval stream stage measurements, using a stage-discharge rating curve calibrated to manual Q measurements made under baseflow and stormflow conditions. Stream stage was measured with non-vented pressure transducers (PT) (Solinst Levelogger or Onset HOBO Water Level Logger) installed in stilling wells at two locations in Faga'alu: FOREST and VILLAGE. Stream gauging sites were chosen to take advantage of existing control structures (FOREST) or stabilized stream cross sections (VILLAGE). Barometric pressure data to calculate stage from PT's was collected at Wx. Data gaps were filled by barometric data from stations at Pago Pago Harbor (NSTP6) and NOAA Climate Observatory at Tula (TULA). Priority was given to the station closest to the watershed with valid barometric pressure data.")
document.add_paragraph("Discharge (Q) was measured in the field by the area-velocity method (AV) using a Marsh-McBirney flowmeter to measure flow velocity and simultaneous channel surveys to measure cross-section geometry (Harrelson et al., 1994; Turnipseed and Sauer, 2010). AV measurements were made at FOREST and VILLAGE in Faga'alu, and linear, log-linear, and nonlinear rating curves were tested for best fit. AV measurements could not be made at high stages at FOREST and VILLAGE for safety reasons, and peak stages were more than double the highest manual stage-Q measurements. The stage-discharge rating at VILLAGE was extrapolated using Manning's equation by calibrating Manning's n to the manual AV measurments. The flow control structure at FOREST does not meet the assumptions for modeling Q with Manning's equation so the rating at FOREST was modeled with HEC-RAS (Brunner, 2010) and calibrated to the manual AV measurements. ")
## Suspended Sediment Yield
document.add_heading('Suspended Sediment Concentration and Turbidity',level=4)
def No_All_samples(location):
    No_samples = len(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin([location])])
    return No_samples
document.add_paragraph("Stream water samples for SSC were collected by grab or "+'"dip"'+" sampling with 500mL HDPE bottles at FOREST, QUARRY, and VILLAGE, and by an ISCO 3700 Autosampler at QUARRY. From January 6, 2012, to October 1, 2014, "+str(len(SSC_dict['Pre-ALL']))+" samples were collected at 12 sites in Faga'alu and analyzed for SSC. Samples were analyzed in the field using gravimetric methods (Gray et al., 2000). Water samples were vacuum filtered on 47mm, 0.7um Millipore AP40 glass fiber filters, oven dried at 100C for one hour, cooled and weighed to determine SSC (mg/L). Other sites were sampled opportunistically during storm events, but the three main sampling locations in Faga'alu are the focus in this analysis: 1) Upstream (FOREST)(n="+"%.0f"%No_All_samples('DAM')+"), 2) immediately downstream of the Quarry (QUARRY)(n="+"%.0f"%No_All_samples('DT')+" from grab sampling, n="+"%.0f"%No_All_samples('R2')+" by the Autosampler), and 3) Downstream (VILLAGE)(n="+"%.0f"%No_All_samples('LBJ')+").")
document.add_paragraph("Suspended sediment concentration (SSC) at 15 min intervals was derived from 1) 15Min interval turbidity (T) data, using a T-SSC relationship calibrated to stream water samples collected over a range of Q and SSC, and 2) interpolating SSC from grab and autosamples, provided more than three samples were collected during the storm event, and they adequately captured the peak SSC of the storm event. The T-SSC relationship is unique to each region, or even each stream, and can be influenced by water color, dissolved solids and organic matter, temperature, and the shape, size, and composition of sediment particles. However, T has proved to be a robust surrogate measure of SSC in streams (Gippel 1995) and is widely used for remote monitoring applications (Lewis 1996).")
document.add_paragraph("Turbidity (T) was measured at FOREST and VILLAGE using three types of turbidimeters: 1) a Greenspan TS3000, 2) a YSI 600OMS, and 3) a CampbellSci OBS500. All turbidimeters were permanently installed in protective PVC housings near the streambed where the turbidity probe would be continuously submerged and oriented downstream. Despite regular maintenance, debris fouling during storm and baseflows was common and caused data loss during several storm events. The YSI turbidimeter was recalibrated with YSI NTU standards at the beginning of each successive field season, approximately every 3-6 months during data collection. All turbidimeters were regularly cleaned following storms to ensure proper operation.")
document.add_paragraph("At FOREST, a Greenspan TS3000 turbidimeter recorded T (NTU) at 5Min intervals from January 2012 until it was vandalized and destroyed in July 2012. The YSI turbidimeter previously deployed at VILLAGE in 2012 was repaired by YSI and redeployed at FOREST and recorded T (NTU) at 5Min intervals from June 2013 to October 2013, and January 2014 to August 2014. Turbidity data was resampled to 15 min intervals to compare with SSC samples for the T-SSC relationship, and corresponding to Q for calculating suspended sediment yield (SSY) (Equation "+SSYEV_eq.eq_num+").")
document.add_paragraph("At VILLAGE, a YSI 600OMS sonde with 6136 Turbidity Probe recorded T (NTU) at 5Min intervals from January 30, 2012, to February 20, 2012, and at 15Min intervals from February 27, 2012 to May 23, 2012, when it was damaged during a large storm. The YSI turbidimeter was replaced with a CampbellSci OBS500 which recorded Backscatter (BS) and Sidescatter (SS, comparable to NTU) at 5Min intervals from March 7, 2013, to July 15, 2014. No data was recorded from August 2013-January 2014 due to instrument malfunction (wiper clogged with sediment). A new CampbellSci OBS500 was installed at VILLAGE from January, 2014, to August, 2014. To correct for some periods of high scatter observed in the BS and SS data recorded by the CampbellSci OBS500 in 2013, the new CampbellSci OBS500 was programmed to make 100 BS and SS measurements every 15Min, and record Median, Mean, STD, Min, and Max BS and SS.")
document.add_paragraph("It is assumed that sediment sources change during storm periods, and at different points in the watershed, so a unique T-SSC relationship was developed for each turbidimeter at each location using 15Min interval T data and SSC samples from storm periods only. A "+'"synthetic"'+" T-SSC relationship (SRC) was also developed by placing the turbidimeter in a black tub, and sampling T and SSC as sediment was added. An SRC was developed for the YSI 600OMS turbidimeter at FOREST and the CampbellSci OBS500 at VILLAGE, in 2014, using sediment collected from the streambed near the turbidimeter. An SRC was not developed for the TS300 turbidimeter at FOREST, or the YSI 600OMS turbidimeter when it was installed at VILLAGE in 2012. ")

## Measurement Uncertainty
document.add_heading('Measurement Uncertainty',level=3)
document.add_paragraph("Uncertainty in SSYEV estimates arises from both measurement and model errors, including models of stage-discharge (stage-Q) and turbidity-suspended sediment concentration (T-SSC) (Harmel et al., 2006). The Root Mean Square Error (RMSE) method estimates the "+'"most probable value"'+" of the cumulative or combined error by propagating the error from each measurement and modeling procedure to the final SSYEV estimate (Topping, 1972). The resulting cumulative probable error (PE) is the square root of the sum of the squares of the maximum values of the separate errors:")
add_equation(PE_eq) ## Equation
document.add_paragraph("Error from manual water discharge measurements using the Area-Velocity method, from continuous discharge measurment in a natural channel, from grab sampling and autosampling SSC during stormflows , and from lab procedures are considered "+'"measurement errors"'+" and were estimated using lookup tables from the DUET-H/WQ software tool (Harmel et al., 2006). These measurement errors (RMSE) were combined with the modeling errors (RMSE) from the stage-Q and T-SSC relationships to calculate PE for each storm event, to add a statistical measure of uncertainty to SSYEV (plus-minus tons). The effect of uncertain SSYEV estimates may complicate conclusions about SSYEV-Qmax relationships, contributions from subwatersheds, and anthropogenic impacts. This is common in sediment yield studies where successful models estimate SSY with plus-minus 50-100% accuracy (Duvert et al., 2012), but preliminary data suggested the difference in SSYEV from the upper and lower subwatersheds is significantly larger than the ranges of uncertainty in the SSY estimates. ")

#### RESULTS ####
results_title = document.add_heading('Results',level=2)
## Field data collection
document.add_heading('Field Data Collection',level=3)

#### Precipitation
document.add_heading('Precipitation',level=4)
document.add_paragraph("Annual precipitation measured at RG1 was "+"%.0f"%PrecipFilled[start2012:stop2012].sum()+"mm, "+"%.0f"%PrecipFilled[start2013:stop2013].sum()+"mm,  and "+"%.0f"%PrecipFilled[start2014:stop2014].sum()+" in 2012, 2013, and 2014 respectively. These annual rainfall amounts are approximately 75% of long-term rainfall data (=4500-4800mm) from the Parameter-elevation Relationships on Independent Slopes Model (PRISM) (Craig, 2009). Comparison of rain gauge data showed no orographic relationship between RG1 and Wx, or RG1 and RG2, so precipitation was assumed to be homogenous over the watershed for all analyses. Rain gauges could only be placed as high as ~300m (RG2), though the highest point in the watershed is ~600m. Long-term rain gage records show a strong precipitation gradient with increasing elevation, with average precipitation of 3-4,000mm on the lowlands, increasing to more than 6,350mm at the highest elevations around the harbor (Mt. Alava and Matafao Peak) (Craig, 2009; Dames & Moore, 1981; Wong, 1996). Rainfall data measured at higher elevations would be useful to determine a more robust orographic rainfall relationship. For this analysis, however, precipitation is only used as a predictive storm metric, and the absolute values of total rainfall in each subwatershed are not as important.")

#### Water Discharge
document.add_heading('Water Discharge',level=4)
document.add_paragraph("At FOREST, recorded stage varied from "+"%.0f"%DAM['stage'].min()+" to "+"%.0f"%DAM['stage'].max()+"cm. Manual discharge measurements (n= "+"%.0f"%len(DAMstageDischarge)+") were made from "+"%.0f"%DAMstageDischarge['Q-AV(L/sec)'].min()+" to "+"%.0f"%DAMstageDischarge['Q-AV(L/sec)'].max()+" L/sec, covering a range of stages from "+"%.0f"%DAMstageDischarge['stage(cm)'].min()+" to "+"%.0f"%DAMstageDischarge['stage(cm)'].max()+"cm. Since the highest recorded stage was more than 2x the highest stage with measured Q, the rating could not be extrapolated by mathematical methods like a power law, so the HEC-RAS model was used. The surveyed geometry of the upstream channel and flow structure at FOREST was input to HEC-RAS, and the HEC-RAS model was calibrated to the manual Q measurements (Figure "+DAM_StageDischarge['fig_num']+"). The RMSE for the HEC-RAS model and Q measurements was "+"%.0f"%DAM_HEC_rmse+" (L/sec), which is used to calculate the total Probable Error.")

DAM_Stormflow_conditions = DAM[DAM['stage']==DAM_storm_threshold.round(0)]['Q'][0]
LBJ_Stormflow_conditions = LBJ[LBJ['stage']==LBJ_storm_threshold.round(0)]['Q'][0]

document.add_paragraph("At VILLAGE, recorded stage varied from "+"%.0f"%LBJ['stage'].min()+" to "+"%.0f"%LBJ['stage'].max()+"cm. Manual discharge measurements (n= "+"%.0f"%len(LBJstageDischarge)+") were made from "+"%.0f"%LBJstageDischarge['Q-AV(L/sec)'].min()+" to "+"%.0f"%LBJstageDischarge['Q-AV(L/sec)'].max()+" L/sec, covering a range of stages from "+"%.0f"%LBJstageDischarge['stage(cm)'].min()+" to "+"%.0f"%LBJstageDischarge['stage(cm)'].max()+"cm. Since the highest recorded stage was more than 2x the highest stage with measured Q, the rating could not be extrapolated by mathematical methods, so the rating was extrapolated using Manning's equation and surveyed stream geometry. Manning's n parameter was calibrated using the manual Q measurements (Figure "+LBJ_StageDischarge['fig_num']+"). The RMSE for Mannings modeled Q and Q measurements was "+"%.0f"%LBJ_Man_rmse+" (L/sec), which is used to calculate the total Probable Error.")
document.add_paragraph("Discharge at both FOREST and VILLAGE was characterized by periods of low but perennial baseflow (FOREST: "+"%.0f"%DAM['Q'].min()+"-"+"%.0f"%DAM_Stormflow_conditions+" L/sec; VILLAGE: "+"%.0f"%LBJ['Q'].min()+"-"+"%.0f"%LBJ_Stormflow_conditions+" L/sec), punctuated by short, flashy hydrograph peaks (FOREST: max "+"%.0f"%DAM['Q'].max()+" L/sec, VILLAGE: max "+"%.0f"%LBJ['Q'].max()+" L/sec). ")

## DAM stage-discharge rating
if 'DAM_StageDischarge' in locals():
    document.add_picture(DAM_StageDischarge['filename']+'.png',width=Inches(6))
    add_figure_caption(DAM_StageDischarge['fig_num'],"Stage-Discharge relationships for stream gaging site at FOREST.")
## LBJ stage-discharge rating
if 'LBJ_StageDischarge' in locals():
    document.add_picture(LBJ_StageDischarge['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_StageDischarge['fig_num'],"Stage-Discharge relationships for stream gaging site at VILLAGE.")
## Storm Water Discharge
if 'Q_Diff_table' in locals():
    dataframe_to_table(df=Q_Diff_table,table_num=Q_Diff_table.table_num,caption="Water discharge from subwatersheds in Faga'alu")

#### Suspended Sediment Concentration
document.add_heading('Suspended Sediment Concentration',level=4)
## Mean and Max SSC  numbers
Mean_SSC_Forest, Max_SSC_Forest  = SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].mean(), SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].max()
Mean_SSC_Quarry, Max_SSC_Quarry = SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DT'])]['SSC (mg/L)'].mean(), SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DT'])]['SSC (mg/L)'].max()
Mean_SSC_Village, Max_SSC_Village = SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].mean(), SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].max()
Max_SSC_Village_Event = "{:%m/%d/%Y}".format(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax())
Max_SSC_Village_Q = LBJ['Q'][SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['LBJ'])]['SSC (mg/L)'].idxmax()]
Max_SSC_Forest_Event = "{:%m/%d/%Y}".format(SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax())
Max_SSC_Forest_Q = DAM['Q'][SSC_dict['Pre-ALL'][SSC_dict['Pre-ALL']['Location'].isin(['DAM'])]['SSC (mg/L)'].idxmax()]

document.add_paragraph("Mean SSC of grab samples collected during baseflow and stormflow were lowest at the FOREST site ("+"%.0f"%Mean_SSC_Forest+" mg/L), and highest at the QUARRY ("+"%.0f"%Mean_SSC_Quarry+" mg/L) and the donwstream VILLAGE site ("+"%.0f"%Mean_SSC_Village+" mg/L) . Maximum SSC values were "+"%.0f"%Max_SSC_Forest+" mg/L, "+"%.0f"%Max_SSC_Quarry+" mg/L, and "+"%.0f"%Max_SSC_Village+" mg/L at the upstream (FOREST), quarry (QUARRY), and downstream (VILLAGE) sites, respectively. The maximum SSC values at QUARRY ("+"%.0f"%Max_SSC_Quarry+" mg/L ) and VILLAGE ("+"%.0f"%Max_SSC_Village+" mg/L) sites were sampled during the same event "+Max_SSC_Village_Event+", during fairly low discharge (Q_VILLAGE="+"%.0f"%Max_SSC_Village_Q+" L/sec). The maximum SSC value for the upstream site ("+"%.0f"%Max_SSC_Forest+" mg/L) was sampled on "+Max_SSC_Forest_Event+" at high discharge (Q_FOREST= "+"%.0f"%Max_SSC_Forest_Q +" L/sec). Anecdotal and field observations reported higher than normal SSC upstream of the quarry during the 2013 field season (G. Poysky, pers. comm.).")

def No_storm_samples(location):
    No_storm = len(SSC_dict['Pre-storm'][SSC_dict['Pre-storm']['Location'].isin([location])])
    return No_storm
# Percent storm samples in another script, wouldn't work here
def Mean_storm_SSC(location):    
    Mean_storm_samples  = SSC_dict['Pre-storm'][SSC_dict['Pre-storm']['Location'].isin([location])]['SSC (mg/L)'].mean()
    return Mean_storm_samples 

document.add_paragraph("At FOREST, "+"%.0f"%No_storm_samples('DAM')+" grab samples ("+"%.0f"%Percent_storm_Forest+"%) were taken during stormflow conditions (Q_DAM>"+"%.0f"%DAM_Stormflow_conditions+" L/sec), with mean SSC of "+"%.0f"%Mean_storm_SSC('DAM')+" mg/L. At QUARRY, "+"%.0f"%No_storm_samples('DT')+" grab samples ("+"%.0f"%Percent_storm_Quarry +"%) were taken during stormflow conditions (Q_DAM>"+"%.0f"%DAM_Stormflow_conditions+" L/sec); and mean SSC for QUARRY stormflow samples was "+"%.0f"%Mean_storm_SSC('DT')+" mg/L. At VILLAGE, "+"%.0f"%No_storm_samples('LBJ')+" samples ("+"%.0f"%Percent_storm_Village+"%) were taken during stormflow conditions (Q_VILLAGE>"+"%.0f"%LBJ_Stormflow_conditions+" L/sec); and mean SSC for VILLAGE stormflow samples was "+"%.0f"%Mean_storm_SSC('LBJ')+" mg/L (Figure "+SSC_Boxplots['fig_num']+". This pattern of SSC values suggests there is a large input of sediment downstream of FOREST at the quarry, and then SSC is diluted by addition of stormflow with lower SSC between QUARRY and VILLAGE.")
## SSC boxplots
if 'SSC_Boxplots' in locals():
    document.add_picture(SSC_Boxplots['filename']+'.png',width=Inches(6))
    add_figure_caption(SSC_Boxplots['fig_num'],"Boxplots of Suspended Sediment Concentration (SSC) from grab samples only (no Autosampler) at FOREST, QUARRY, and VILLAGE during storm periods.")
## Water Discharge vs Sediment Concentration
document.add_paragraph("The highest SSC values were observed just downstream of the quarry, mainly during low or baseflows. Given the close proximity of the quarry to the stream, SSC downstream of the quarry can be highly influenced by mining activity like rock extraction, crushing or hauling operations. Also, prior to 2013, a common practice for removing fine sediment from crushed aggregate was to rinse it with water pumped from the stream. In the absence of retention structures, the fine sediment was discharged directly into the stream, causing high SSC during low Q. While sheetwash erosion of the quarry during storm events causes higher total sediment yield, the instantaneous SSC values are lower during stormflows due to dilution by the water discharged by surrounding forest areas. The practice of manually rinsing fine sediment from aggregate was discontinued in 2013, corresponding with a lack of high SSC grab samples during low discharges (Figure "+Discharge_Concentration['fig_num']+").")
## Discharge Concentration
if 'Discharge_Concentration' in locals():
    document.add_picture(Discharge_Concentration['filename']+'.png',width=Inches(6))
    add_figure_caption(Discharge_Concentration['fig_num'],"Water Discharge vs Suspended Sediment concentration at FOREST, QUARRY, and VILLAGE during baseflow and stormflow periods. Samples from Autosamplerare included with grab samples at QUARRY.")
document.add_paragraph("No clear relationship between Q and SSC could be determined for any of the three sites (FOREST, QUARRY, VILLAGE) due to significant hysteresis observed during storm periods, and the changing sediment availabiltiy associated with the quarrying operations. At FOREST, variability of SSC samples during stormflows from year to year was assumed to be caused by randomly occurring landslides caused by large storms. ")


#### Turbidity
document.add_heading('Turbidity', level=4)
document.add_paragraph(" A "+'"synthetic"'+" T-SSC relationship (SRC) was developed for the YSI 600OMS turbidimeter at FOREST and the CampbellSci OBS500 at VILLAGE, in 2014, using sediment collected from the streambed near the turbidimeter (Figure "+Synthetic_Rating_Curve['fig_num']+"). These relationships were compared with the T-SSC relationships developed using in situ data collected under storm conditions (Figures "+LBJ_and_DAM_YSI_Rating_Curve['fig_num']+"-"+LBJ_OBSb_Rating_Curve['fig_num']+"). In all instances, the SRC's had a steeper slope than the T-SSC relationships from storm samples. However the T-SSC relationships all showed acceptable r2 values ("+"%.2f"%LBJ_YSI_rating.r2+"-"+"%.2f"%DAM_YSI_rating.r2+"), so the SRC's were not used to model SSC from T data.")
document.add_paragraph("The T-SSC relationship varied among sampling sites and sensors. Lower scatter in the linear relationships was achieved by using grab samples collected during stormflows only. It is assumed that the color, particle sizes, and composition of sediment changes during stormflows where much lighter color, finer sediment from the quarry is present. For the TS3000 deployed at FOREST, the r2 value was fairly high ("+"%.2f"%DAM_TS3K_rating.r2+") but the ranges of T and SSC values used to develop the relationship were considered too small to develop a robust relationship for higher T values. Instead, the T-SSC relationship developed for the YSI turbidimeter installed at FOREST was used to convert T data from the TS3000 to SSC. For the YSI 600OMS turbidimeter, more scatter was observed in the T-SSC relationship at VILLAGE than at FOREST, but this could be attributed to the higher number and wider range of values sampled (Figure "+LBJ_and_DAM_YSI_Rating_Curve['fig_num']+"), as well the contribution of multiple sediment sources present at VILLAGE. The CampbellSci OBS500 measures both BS and SS, and both the BS-SSC and SS-SSC relationships at VILLAGE showed high r2 values (Figure "+LBJ_OBSa_Rating_Curve['fig_num']+" and Figure "+LBJ_OBSb_Rating_Curve['fig_num']+"). Mean SS was used since it is physically measured the same as NTU measured by the YSI turbidimeter (Chauncey W. Anderson 2005). The RMSE for each selected T-SSC relationship was computed and used in the estimate of Probable Error. The RMSE for the T-SSC relationship was "+"%.0f"%DAM_YSI_rating.rmse+", "+"%.0f"%LBJ_YSI_rating.rmse+", "+"%.0f"%LBJ_OBSa_rating.rmse+", "+"%.0f"%LBJ_OBSb_rating.rmse+" mg/L for the YSI at FOREST, YSI at VILLAGE, CampbellSci OBS500 at VILLAGE (2013), and CampbellSci OBS500 at VILLAGE (2014), respectively. The RMSE for YSI at FOREST was also used for the TS3000 at FOREST since the same T-SSC relationship was used.")

## Synthetic Rating Curves
if 'Synthetic_Rating_Curve' in locals():
    document.add_picture(Synthetic_Rating_Curve['filename']+'.png',width=Inches(6)) ## add pic from filename defined above
    add_figure_caption(Synthetic_Rating_Curve['fig_num'],"Synthetic Rating Curves for Turbidimeters deployed at FOREST, QUARRY, VILLAGE, Nuuuli-1 and Nuuuli-2")

## LBJ and DAM YSI T-SSC rating curves
if 'LBJ_and_DAM_YSI_Rating_Curve' in locals():
    document.add_picture(LBJ_and_DAM_YSI_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_and_DAM_YSI_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the YSI turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_YSI.dropna().index[-1])+") and at FOREST ("+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(DAM_YSI.dropna().index[-1])+").")

## LBJ OBSa T-SSC rating curve
if 'LBJ_OBSa_Rating_Curve' in locals():
    document.add_picture(LBJ_OBSa_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_OBSa_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the OBS500 turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSa.dropna().index[-1])+").")

## LBJ OBSb T-SSC rating curve
if 'LBJ_OBSb_Rating_Curve' in locals():
    document.add_picture(LBJ_OBSb_Rating_Curve['filename']+'.png',width=Inches(6))
    add_figure_caption(LBJ_OBSb_Rating_Curve['fig_num'],"Turbidity-Suspended Sediment Concentration relationships for the OBS500 turbidimeter deployed at VILLAGE ("+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[0])+"-"+"{:%m/%d/%Y}".format(LBJ_OBSb.dropna().index[-1])+").")



#### Storm Events
document.add_heading('Storm Events',level=3)
No_of_Storm_Intervals = len(LBJ_StormIntervals[LBJ_StormIntervals['start']<Mitigation])
No_of_Storm_Intervals_DAM_Q = len(SedFluxStorms_DAM[SedFluxStorms_DAM['Qstart']<Mitigation]['Qsum'].dropna())
No_of_Storm_Intervals_LBJ_Q = len(SedFluxStorms_LBJ[SedFluxStorms_LBJ['Qstart']<Mitigation]['Qsum'].dropna())
No_of_Storm_Intervals_DAM_S = len(SedFluxStorms_DAM[SedFluxStorms_DAM['Sstart']<Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_LBJ_S = len(SedFluxStorms_LBJ[SedFluxStorms_LBJ['Sstart']<Mitigation]['Ssum'].dropna())
No_of_Storm_Intervals_QUA_S = len(S_storm_diff_table_quarry())-1
max_storm_duration = LBJ_StormIntervals['duration (hrs)'].max()/24

document.add_paragraph("Most events showed a typical pattern, where a short period of intense rainfall causes a rapid increase in SSC below the QUARRY, indicating rapid sheetwash of sediment into the stream. Highest event SSC was typically observed at QUARRY, with slightly lower and later peak SSC observed at VILLAGE. SSC at FOREST typically increased slower, remained much lower, and peaked later than the downstream sites. Though peak SSC was highest at QUARRY, the total SSY was highest at VILLAGE due to the addition of storm runoff from the relatively larger watershed downstream. Storm flow at VILLAGE includes both storm runoff from disturbed areas of the quarry and village, and the undisturbed forest on the steep hillsides in the lower watershed (Figure "+Example_Storm['fig_num']+"). Complex storm events occurred when a subsequent period of rainfall followed before Q and SSC fell back to baseflow levels. Several complex events were separated into individual storm events, if they were separated by a reasonable period of time and the stage was nearly at baseflow.")
if 'Example_Storm' in locals():
    document.add_picture(Example_Storm['filename']+'.png',width=Inches(6))
    add_figure_caption(Example_Storm['fig_num'],"Example of storm event ("+"{:%m/%d/%Y}".format(LBJ_StormIntervals.loc[63]['start'])+"). SSY at FOREST and VILLAGE calculated from SSC modeled from T, and SSY at QUARRY from SSC samples collected by the Autosampler.")
    
document.add_paragraph("Using the stage threshold method and manual separation of complex storm events, "+"%.0f"%No_of_Storm_Intervals+" storm events were identified from Q data at VILLAGE from January, 2012, to July 2014. Valid Q data was recorded during "+"%.0f"%No_of_Storm_Intervals_DAM_Q+" events at FOREST (upstream site), and "+"%.0f"%No_of_Storm_Intervals_LBJ_Q+" events at VILLAGE (downstream site). Valid SSC data from T and Interpolated Grab samples was recorded during "+"%.0f"%No_of_Storm_Intervals_DAM_S+" events at FOREST, and "+"%.0f"%No_of_Storm_Intervals_LBJ_S+" events at VILLAGE. Of those storms, "+"%.0f"%len(S_Diff_table)+" events had valid P and SSY data for both the FOREST and VILLAGE to calculate and compare SSY from the UPPER and LOWER watersheds (Table "+S_Diff_table.table_num+"). Valid SSY data from Interpolated grab samples was collected at QUARRY for "+"%.0f"%No_of_Storm_Intervals_QUA_S+" storms to compare with SSY from FOREST and VILLAGE directly (Table "+S_Diff_table_quarry.table_num+"). Storm event durations ranged from "+"%.0f"%LBJ_StormIntervals['duration (hrs)'].min()+" hours to "+"%.0f"%max_storm_duration+" days, with mean duration of "+"%.0f"%LBJ_StormIntervals['duration (hrs)'].mean()+" hours.") 


#### Comparing SSY from disturbed and undisturbed subwatersheds
document.add_heading('Comparing SSY from disturbed and undisturbed subwatersheds',level=3)
document.add_paragraph("Two approaches were used to determine relative contributions to total SSY from the undisturbed and human-disturbed areas: comparing individual event and average percent contributions, and comparing event-based storm metrics and area-normalized SSYEV.")
Percent_Upper_S = S_storm_diff_table()['% Upper'][-1]
Percent_Lower_S = S_storm_diff_table()['% Lower'][-1]
document.add_paragraph("Percent contributions were calculated to determine relative SSYEV contributions by the upper, undisturbed subwatershed (UPPER), and the lower, human-disturbed subwatershed (LOWER) for storm events with available data at both locations (FOREST and VILLAGE). SSY from the LOWER watershed (SSYLower), including the quarry, was calculated by subtracting SSY measured at FOREST from SSY measured at VILLAGE (SSYLower = SSYVILLAGE - SSYFOREST). SSY from the UPPER subwatershed (SSYUPPER) accounted for an average of "+Percent_Upper_S+"% of Total SSY to Fag'alu Bay, while SSY from the Lower watershed (SSYLower) accounted for an average of "+Percent_Lower_S+"% of Total SSY (SSYTotal) (Table "+S_Diff_table.table_num+").")
## Storm Sediment Table
if 'S_Diff_table' in locals():
    dataframe_to_table(df=S_Diff_table,table_num=S_Diff_table.table_num,caption="Sediment discharge from subwatersheds in Faga'alu")
document.add_paragraph('')

Percent_QUARRY_S = S_storm_diff_table_quarry()['% Quarry'][-1]
Percent_Upper_S = S_storm_diff_table_quarry()['% Upper'][-1]
Percent_Lower_S = S_storm_diff_table_quarry()['% Lower'][-1]
Percent_Village_S = float(Percent_Lower_S) - float(Percent_QUARRY_S)
document.add_paragraph("SSYEV data for "+"%.0f"%No_of_Storm_Intervals_QUA_S+" storms measured at QUARRY was compared with SSYLower, to determine the SSYEV contribution from just the quarry. For the measured storms, which represent a small fraction of storms measured at the FOREST and VILLAGE locations, "+Percent_Upper_S+"% of Total SSY came from the Upper subwatershed, "+Percent_Lower_S+"% of Total SSY was from the Lower subwatershed, and "+Percent_QUARRY_S+"% of Total SSY was from just the quarry (included in the Lower subwatershed). This suggests that "+"%.0f"%Percent_Village_S+"% of Total SSY is from just the village and forest areas in the Lower subwatershed downstream of the quarry, which is comparable to the contribution from the simlarly sized Upper watershed.")

if 'S_Diff_table_quarry' in locals():
    dataframe_to_table(df=S_Diff_table_quarry,table_num=S_Diff_table_quarry.table_num,caption="Sediment discharge from subwatersheds in Faga'alu")
#document.add_paragraph("Storm on 2/3/12 has a potential outlier at DT at beginning of storm, which makes the SSYquarry huge! Storm at 2/5/12 doesn't have adequate SSC samples for quarry and its a multipeaked event so the SSY doesn't fall back to low levels like the LBJ and DAM T data suggest it should. Storm 3/6/13 looks like may have missed the second peak but the data looks comparable between sites. Storm 4/23/13 looks good. Storm 4/30/13 has inadequate SSC data for all locations after the first peak; can maybe change the storm interval? Storm 6/5/13 has inadequate SSC data for all locations for the first peak, decent data for LBJ and DT for second peak but not for DAM; can maybe change the storm interval? Storm 2/14/14 looks good. Storm 2/20/14 looks good. Storm 2/21/14 looks good. Storm 2/27/14 looks kinds shitty. So good storms are: 3/6/13, 4/23/13, 2/14/14, 2/20/14, 2/21/14") 
document.add_paragraph("")
if 'Q_S_Diff_summary_table' in locals():
    dataframe_to_table(df=Q_S_Diff_summary_table,table_num=Q_S_Diff_summary_table.table_num,caption="Total Q and SSY")

## Disturbance Ratio
SSYspec_Forest = Q_S_Diff_summary_table.ix['SSY* Forest'][''][:4]
SSYspec_Village = Q_S_Diff_summary_table.ix['SSY* Village'][''][:4]
DR_S = Q_S_Diff_summary_table.ix['SSY* Village'][''][:4]
DR_Q = Q_S_Diff_summary_table.ix['Q Disturbance Ratio'][''][:4]
document.add_paragraph("The disturbance ratio (DR) is the ratio of SSY from the total human-disturbed watershed under current conditions (measured at the watershed outlet-VILLAGE) to SSY under pre-disturbance conditions (SSYpre)(Equation "+DR_eq.eq_num +"). The specific SSY for both FOREST and VILLAGE were calculated by summing all storms from Table ("+S_Diff_table.table_num+") and dividing by the subwatershed area. The DR for water discharge was also calculated to determine if observed changes in SSY were attributable to errors in quantifying Q. The DR for water discharge was "+DR_Q+", suggesting that specific water discharge from the subwatersheds is the same, which is expected since they are similar sizes. The DR for SSY was "+DR_S+", meaning human disturbance, mainly from the quarry has increased Total SSY "+DR_S+"x over undisturbed levels.")  

#### Fitting SSY models
document.add_heading('Fitting SSYEV models',level=3)
document.add_paragraph("The other approach to quantify the impact of human-disturbance in Faga'alu was to compare the specific SSY from the undisturbed and the total watershed using storm metrics. Four storm metrics were assessed: Total Precipitation (Psum), Erosivity Index (EI30), Total Water Discharge (Qsum), and Peak Water Discharge (Qmax). Ordinary Least Squares regressions were fit to log-transformed SSY measured at FOREST and VILLAGE normalized by watershed area and storm metric data (Figure "+SSY_models_ALL['fig_num']+").")

if 'SSY_models_ALL' in locals():
    document.add_picture(SSY_models_ALL['filename']+'.png',width=Inches(6))
    add_figure_caption(SSY_models_ALL['fig_num'],"SSY rating curves for predictors")

document.add_paragraph("Coefficients of determination (r2) were calculated for POWER law regressions fit to the SSY-Qmax relationship for UPPER and TOTAL watersheds. The LOWER subwatershed was excluded for this analysis since a suitable method to normalize Qmax for the subwatershed could not be found in the literature. The LINEAR function achieved a better fit (r2 =0.77) than the POWER law function (r2=0.71) for the UPPER watershed, but not for the TOTAL watershed (r2: power="+' Number '+"). A main objective of this study was to develop a model of SSY to Faga'alu Bay so the best fit model for the TOTAL watershed was selected. All examples found in the literature also used the Power law function, so for better comparison with other studies, the POWER law model fit to event SSY-Qmax was selected as the best model (Figure "+SSY_models_ALL['fig_num']+", Table "+Pearson_table.table_num+", Table"+Spearman_table.table_num+"). Future analysis could also compare the log-transformed linear with bias correction and nonlinear fitting methods in Duvert et al. (2012) but they were not performed for this analysis.")


# Pearson's correlation coeffs
#document.add_paragraph("Pearson's correlation coefficients...")
if 'Pearson_table' in locals():
    dataframe_to_table(df=Pearson_table,table_num=Pearson_table.table_num,caption="Pearson correlation coefficients")
# Spearman's correlation coeffs
#document.add_paragraph("Spearman's correlation coefficients...")  
if 'Spearman_table' in locals():
    dataframe_to_table(df=Spearman_table,table_num=Spearman_table.table_num,caption="Spearman correlation coefficients")
    
document.add_paragraph("Similar to Duvert et al. (2012), Pearson's and Spearman's correlation coefficients were highest for Qmax for all three watersheds in Faga'alu. Some predictors were highly correlated for a single watershed but not the others, like Qsum for the UPPER watershed. EI was the least correlated with SSY.")    

document.add_paragraph("No statistically significant correlations (p<0.001) were found between SSY and EI using Pearson's or Spearman's coefficients. Overall, statistically significant Pearson's and Spearman's correlation coefficients were fairly similar, indicating the relationship of SSY with the predictor variables is adequately described by a linear function and that outliers and non-normality in the data did not affect the test. The exception was significantly higher Spearman's correlation coefficient for Psum for the TOTAL watershed Spearman's=0.84 vs. Pearson's=0.70). Pearson correlation coefficients were highest overall for Psum and Qmax, indicating these were significantly correlated with event SSY. Qsum showed high correlation with SSY for the UPPER watershed but not the LOWER or TOTAL watershed.")
document.add_paragraph("The Pearson correlation coefficients for the UPPER watershed are slightly higher for discharge-related predictors (especially Qsum) and lower for precipitation-type predictors (Psum and EI30) than for the LOWER watershed. This suggests that sediment production processes are more dominated by discharge in the UPPER watershed and precipitation in the LOWER watershed. SSY from the LOWER watershed is hypothesized to be mostly generated by surface erosion at the quarry, dirt roads, and agricultural plots, whereas SSY from the UPPER watershed is hypothesized to be mainly from channel processes and mass wasting. Mass wasting can contribute large pulses of sediment during large precipitation events, which can be deposited in lobes on the streambanks and entrained at high discharges during later events. Qmax may be a promising predictor that integrates these processes.")


document.add_paragraph("Qmax was the most correlated with SSYEV for all three watersheds, and was selected as the best predictor variable . To investigate for any patterns not reflected in the Pearson's and Spearman's correlation coefficients, regression lines were fit to all four predictor variables using ordinary least squares regression on log-transformed data (POWER law)(Figure "+SSY_models_ALL['fig_num']+"). ")
document.add_paragraph("Interpret Figure 11 for the reader. What do you want them to see in Figure 11? What is important?")

  
#### DISCUSSION
discussion_title=document.add_heading('Discussion',level=2)
document.add_paragraph("EI30 is poorly correlated with SSY due to the effect of previous events (Rodrigues et al., 2013)")

## Hypotheses
document.add_paragraph("When normalized by area, the SSY-Storm Metric relationship is displaced upward if SSY is higher for the same storm size, which is attributed to the additional sediment loading from human-disturbed areas in the quarry and village. If area-normalized event SSY rating curves converge at higher values, it indicates diminishing human disturbance for large storms. It was hypothesized that at high water discharges, SSY from the Upper watershed may become the dominant source of total sediment loading to Faga'alu Bay.")

document.add_paragraph("Several researchers have attempted to explain the difference in a (intercept) and B (slope) coefficients according to watershed characteristics. A sediment rating curve (Q-SSC) is considered a 'black box' model, and the a and B coefficients have no physical meaning. However, some physical interpretation has been ascribed to them, with the a coefficient representing an erosion severity index, and the B coefficient representing the erosive power of the river. High a values suggest high availability of easily eroded sediment sources in the watershed, and high B values suggest that a small change in stream discharge leads to a large increase in sediment load due to the erosive power of the river or the extent that new sediment sources become available as discharge increases (Asselman, 2000).")

document.add_paragraph("Similar analysis has been done on event-based sediment yield curves. Rankl (2004) found that B coefficients were not statistically different , and he assumed that the B exponent was a function of rainfall intensity on hillslopes. Rankl (2004) hypothesized that variability in a (the intercept) was a function of sediment availability and erodibility in watersheds, but Duvert et al. (2012) argued that a values are dependent on the regression fitting method (nonlinear method fits higher up on low discharges than linear fit).")

## Compare to other areas
document.add_paragraph("The area-normalized SSY-Qmax models for the UPPER and TOTAL watersheds in Faga'alu were compared to the SSY-Qmax models from eight small, mountainous catchments in Spain, France and Mexico (Duvert et al., 2012). The UPPER watershed model was slightly higher, and the TOTAL watershed model was significantly higher than the nonlinear model fit to all the data in Duvert et al. (2012)(Figure 13. The UPPER model was very similar to a watershed specific model (ARES) for a watershed in the Spanish Pyrenees (0.45km2).")

## Introduce the post-mitigation work
document.add_paragraph("Sediment runoff management was implemented at the quarry October 1, 2014. Storm monitoring is currently underway and results will be presented soon...")

#### CONCLUSION
conclusion_title=document.add_heading('Conclusion',level=2)
conclusion_text = document.add_paragraph('Conclusion text goes here...')


## 
## Save Document
document.save(maindir+'Manuscript/DRAFT.docx')

## Clean up any open figures
plt.close('all')








